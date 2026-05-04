#!/usr/bin/env python3
"""
Main extraction pipeline for doc2ml-json.

Converts any supported document (PDF, EPUB, DOCX, TXT, MD, HTML) into
structured JSON following the doc2ml-json schema v0.6.2.
"""

import argparse
import hashlib
import json
import os
import re
import sys
import time
import unicodedata
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import parse_qs, urlencode, urlparse, urlunparse

# Optional third-party libraries with graceful fallbacks
try:
    import fitz
except ImportError:
    fitz = None
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    from bs4 import BeautifulSoup, NavigableString, Tag
except ImportError:
    BeautifulSoup = None
try:
    from docx import Document
except ImportError:
    Document = None
try:
    import yaml
except ImportError:
    yaml = None
try:
    from lxml import etree
except ImportError:
    etree = None
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


def sha256_file(filepath: str) -> str:
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def _serialize_dates(obj):
    """Recursively convert datetime/date objects to ISO strings for JSON serialization."""
    from datetime import date, datetime
    if isinstance(obj, dict):
        return {k: _serialize_dates(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_serialize_dates(v) for v in obj]
    if isinstance(obj, datetime):
        return obj.isoformat()
    if isinstance(obj, date):
        return obj.isoformat()
    return obj


def normalize_unicode(text: str) -> str:
    return unicodedata.normalize("NFKC", text)


def normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    for ch in "\u200b\u200c\u200d\ufeff\u2060\u00ad":
        text = text.replace(ch, "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(lines)


def estimate_tokens(text: str) -> int:
    return max(1, int(len(text.split()) * 0.75))


EXCLUDED_KEYS = {"wikipedia", "see", "arxiv", "http", "https", "doi", "ref", "source"}


def extract_key_values(text: str) -> dict:
    """Heuristic key-value extraction for patterns like 'Key: Value' or 'Key → Value'.
    Returns a dict only when >1 lines in the text match the pattern."""
    kvs = {}
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^([^:]+?)\s*[:\u2192]\s*(\S.*)$", line)
        if m:
            key = m.group(1).strip()
            val = m.group(2).strip()
            if not key or not val:
                continue
            key_lower = key.lower()
            if key_lower in EXCLUDED_KEYS:
                continue
            if "wikipedia:" in val.lower() or "arxiv:" in val.lower():
                continue
            if any(kw in key_lower for kw in (".com", ".org", ".net", "http")):
                continue
            kvs[key] = val
    return kvs if len(kvs) > 1 else {}


def _extract_cross_references_from_blocks(blocks: list[dict]) -> list[dict]:
    """Scan all block text for cross-references (URLs, arXiv, DOI, Wikipedia, emails)."""
    patterns = [
        ("url", r'https?://[^\s<>"\')\]]+'),
        ("arxiv", r'ar[Xx]iv:\d{4}\.\d{4,5}(?:v\d+)?'),
        ("doi", r'10\.\d{4,}/[^\s]+'),
        ("wikipedia", r'[Ww]ikipedia:[^\s]+'),
        ("email", r'[\w.-]+@[\w.-]+\.\w+'),
    ]
    refs = []
    seen = set()
    for b in blocks:
        text = b.get("text_plain", "")
        if not text:
            continue
        for ref_type, pattern in patterns:
            for m in re.finditer(pattern, text):
                matched = m.group(0)
                # Trim trailing punctuation likely not part of the reference
                matched = matched.rstrip(".,;:!?)")
                if matched in seen:
                    continue
                seen.add(matched)
                refs.append({
                    "type": ref_type,
                    "target": matched,
                    "context": text[max(0, m.start() - 30):m.end() + 30],
                    "source_chunk_id": b.get("chunk_id"),
                })
    return refs


def _extract_dates_from_text(text: str) -> list[str]:
    """Extract dates from text and return ISO 8601 formatted dates."""
    month_names = "January|February|March|April|May|June|July|August|September|October|November|December"
    pattern = rf'\b({month_names})\s+(\d{{1,2}}),?\s+(\d{{4}})\b'
    matches = re.finditer(pattern, text, re.IGNORECASE)
    dates = []
    month_map = {
        'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
        'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
    }
    for m in matches:
        month_str, day, year = m.group(1), int(m.group(2)), int(m.group(3))
        month_num = month_map[month_str.lower()]
        dates.append(f"{year:04d}-{month_num:02d}-{day:02d}")
    return dates


def clean_url(url: str, base_url: str | None = None) -> str:
    if not url:
        return url
    tracking = ["utm_source", "utm_medium", "utm_campaign", "utm_term", "utm_content", "fbclid", "gclid", "ref", "source"]
    if base_url and not urlparse(url).netloc:
        from urllib.parse import urljoin
        url = urljoin(base_url, url)
    parsed = urlparse(url)
    if parsed.query:
        q = parse_qs(parsed.query)
        for p in tracking:
            q.pop(p, None)
        parsed = parsed._replace(query=urlencode(q, doseq=True))
    if parsed.scheme == "http":
        parsed = parsed._replace(scheme="https")
    return urlunparse(parsed)


def make_block_base(chunk_id: str, block_type: str, text_plain: str, page_num: int | None = None,
                     source_location: str = "", extraction_method: str = "", confidence: float = 0.95) -> dict:
    return {
        "chunk_id": chunk_id,
        "type": block_type,
        "content": {},
        "text_plain": text_plain,
        "text_original": text_plain,
        "char_count": len(text_plain),
        "token_count_est": estimate_tokens(text_plain),
        "embedding_ready": block_type not in ("metadata", "page_header", "page_footer", "footnote", "unknown"),
        "context_window": {
            "prev_chunk_id": None,
            "next_chunk_id": None,
            "parent_heading_chunk_id": None,
            "parent_structure_node_id": None,
            "surrounding_text_preview": "",
        },
        "provenance": {
            "page_number": page_num,
            "page_range": None,
            "bounding_box": None,
            "source_location": source_location,
            "extraction_method": extraction_method,
            "confidence": confidence,
        },
        "language": {"detected": "en", "confidence": 0.99},
        "semantics": {"heading_level": None, "is_first_paragraph": False, "is_last_paragraph": False,
                      "section_role": None, "sentiment_score": None, "readability_flesch": None},
        "relations": {"part_of_table": None, "part_of_list": None, "part_of_figure": None,
                        "footnotes_for": [], "references": []},
        "custom": {},
    }


def detect_format(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    ext_map = {
        ".pdf": "application/pdf", ".epub": "application/epub+zip",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain", ".md": "text/markdown", ".markdown": "text/markdown",
        ".html": "text/html", ".htm": "text/html",
    }
    mime = ext_map.get(ext)
    if mime:
        return mime
    with open(filepath, "rb") as f:
        header = f.read(8)
    if header.startswith(b"%PDF-"):
        return "application/pdf"
    if header.startswith(b"PK"):
        try:
            with zipfile.ZipFile(filepath, "r") as z:
                nl = z.namelist()
                if "word/document.xml" in nl:
                    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if any(".opf" in n for n in nl):
                    return "application/epub+zip"
        except Exception:
            pass
    if b"<html" in header.lower() or b"<!doctype" in header.lower():
        return "text/html"
    return "text/plain"


def _estimate_heading_level_from_text(text: str) -> int:
    """Estimate heading level from text characteristics.

    Returns 1-6 heading level based on formatting heuristics.
    """
    if not text:
        return 2
    if text.isupper() and len(text) < 40:
        return 1
    if text.isupper() and len(text) < 60:
        return 2
    # Numbered patterns: "1." -> level 1, "1.1." -> level 2, "1.1.1" -> level 3
    m = re.match(r'^(\d+(?:\.\d+)*)', text)
    if m:
        depth = m.group(1).count('.') + 1
        return min(depth, 6)
    if len(text.split()) <= 3:
        return 3
    return 2


def _split_page_into_blocks(page_text: str, page_number: int, base_block_idx: int = 0) -> list[dict]:
    """Split a page of text into multiple structured blocks.

    Returns list of dicts with keys: type, text, page_number, and optionally
    level (for headings), ordered/indent (for list_items).
    """
    blocks = []
    lines = page_text.split('\n')
    current_paragraph = []

    def flush_paragraph():
        if current_paragraph:
            text = ' '.join(current_paragraph).strip()
            if text:
                blocks.append({'type': 'paragraph', 'text': text, 'page_number': page_number})
            current_paragraph.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            continue

        # Heading detection
        # Pattern 1: Numbered section ("1. Title", "1.1 Subtitle", "2.1.1 Deep")
        if re.match(r'^\d+(\.\d+)*\.?\s+\S', stripped) and len(stripped) < 120:
            # Exclude time/data conversion patterns ("1 kilosecond: 16.7 minutes", "1 Earth day: 86.4 ks")
            if re.search(r'^\d+\s+(\w+|\w+\s+\w+):\s*\d', stripped):
                pass  # fall through to paragraph
            else:
                flush_paragraph()
                blocks.append({
                    'type': 'heading',
                    'text': stripped,
                    'page_number': page_number,
                    'level': _estimate_heading_level_from_text(stripped),
                })
                continue

        # Pattern 2: Short all-caps line (potential heading)
        if stripped.isupper() and len(stripped) < 80 and len(stripped) > 3:
            # Exclude short labels ending with colon (e.g. "CCAA:")
            if stripped.endswith(':'):
                pass  # fall through to paragraph
            else:
                flush_paragraph()
                blocks.append({
                    'type': 'heading',
                    'text': stripped,
                    'page_number': page_number,
                    'level': _estimate_heading_level_from_text(stripped),
                })
                continue

        # Pattern 3: Short line, no sentence-ending punctuation (potential heading)
        # Must NOT look like a list item (exclude bullets and numbered items)
        if (len(stripped) < 80 and not any(c in stripped for c in '.!?;')
                and len(stripped.split()) <= 8
                and not re.match(r'^[\-\u2022*\u00b7\u25e6\u25cb\d]', stripped)):
            lower = stripped.lower()
            words = stripped.split()
            skip = False
            # Skip citation markers (See:, Source:, Note:)
            if re.match(r'^(See|Source|Note|Ref|Also)\s*:', stripped, re.IGNORECASE):
                skip = True
            elif any(tag in lower for tag in ['wikipedia:', 'arxiv:', 'doi:', 'http://', 'https://']):
                skip = True
            elif lower == 'edit':
                skip = True
            elif stripped.endswith(':') and len(words) <= 3:
                skip = True
            elif len(words) == 1 and not any(c.isalnum() for c in words[0]):
                skip = True
            elif re.match(r'^\d+\s+\w+:\s*\d', stripped):
                skip = True
            elif len(words) >= 2 and words[1].startswith(':') and len(words[0]) > 2:
                skip = True
            if not skip:
                flush_paragraph()
                blocks.append({
                    'type': 'heading',
                    'text': stripped,
                    'page_number': page_number,
                    'level': _estimate_heading_level_from_text(stripped),
                })
                continue

        # List detection
        list_match = re.match(r'^(\s*)([-\u2022*\u00b7\u25e6\u25cb]\s+|\d+[.)]\s+)(.*)', stripped)
        if list_match:
            flush_paragraph()
            blocks.append({
                'type': 'list_item',
                'text': list_match.group(3).strip(),
                'page_number': page_number,
                'ordered': bool(re.match(r'^\d', list_match.group(2).strip())),
                'indent': len(list_match.group(1)),
            })
            continue

        # Glossary / definition list detection
        def_match = re.match(r'^([A-Za-z][A-Za-z\s\-]{1,30}):\s*(.+)', stripped)
        if def_match and len(def_match.group(1).split()) <= 4:
            term = def_match.group(1).strip()
            definition = def_match.group(2).strip()
            if not re.match(r'^\d+\s+\w+', term) and 'http' not in definition[:10]:
                flush_paragraph()
                blocks.append({
                    'type': 'definition_list',
                    'term': term,
                    'definition': definition,
                    'text': stripped,
                    'page_number': page_number,
                })
                continue

        current_paragraph.append(stripped)

    flush_paragraph()

    # Post-process: merge consecutive short headings that form a split title
    merged_blocks = []
    i = 0
    while i < len(blocks):
        blk = blocks[i]
        if (blk['type'] == 'heading' and i + 1 < len(blocks)
                and blocks[i + 1]['type'] == 'heading'
                and len(blk['text']) < 40
                and len(blocks[i + 1]['text']) < 40
                and not any(c in blk['text'][-1] for c in '.!?;')):
            merged_text = blk['text'] + ' ' + blocks[i + 1]['text']
            merged_blocks.append({
                'type': 'heading',
                'text': merged_text,
                'page_number': blk['page_number'],
                'level': blk.get('level', 2),
            })
            i += 2
        else:
            merged_blocks.append(blk)
            i += 1
    return merged_blocks


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------
def extract_pdf(filepath: str) -> dict:
    pages_data = []
    metadata = {}
    if pdfplumber:
        try:
            with pdfplumber.open(filepath) as pdf:
                metadata = pdf.metadata or {}
                for i, page in enumerate(pdf.pages, start=1):
                    raw_text = page.extract_text() or ""
                    words = page.extract_words(keep_blank_chars=False, x_tolerance=3, y_tolerance=3)
                    tables = page.extract_tables() or []
                    pages_data.append({"page_number": i, "raw_text": raw_text, "words": words, "tables": tables})
        except Exception:
            pass
    if fitz and (not pages_data or all(not p["raw_text"].strip() for p in pages_data)):
        try:
            doc = fitz.open(filepath)
            meta = doc.metadata or {}
            metadata.update(meta)
            pages_data = []
            for page in doc:
                text = page.get_text()
                blocks = page.get_text("dict").get("blocks", [])
                pages_data.append({"page_number": page.number + 1, "raw_text": text, "words": [], "tables": [],
                                   "blocks": blocks})
            doc.close()
        except Exception:
            pass
    if not pages_data:
        return {"blocks": [], "metadata": metadata, "pages": 0}

    # Structure inference: split each page into multiple logical blocks
    structured_blocks = []
    for p in pages_data:
        page_text = normalize_whitespace(normalize_unicode(p["raw_text"]))
        if page_text:
            split_blocks = _split_page_into_blocks(page_text, p["page_number"], len(structured_blocks))
            structured_blocks.extend(split_blocks)
        # Add tables as raw blocks for build_structure_and_blocks() to process
        for t in p.get("tables", []):
            if t:
                table_lines = []
                for row in t:
                    cleaned = [str(cell or "").replace("\n", " ").strip() for cell in row]
                    table_lines.append(" | ".join(cleaned))
                table_text = "\n".join(table_lines)
                structured_blocks.append({
                    "type": "table",
                    "rows": t,
                    "text": table_text,
                    "page_number": p["page_number"],
                })

    # Build raw blocks from structured blocks for downstream processing
    raw_blocks = []
    for i, blk in enumerate(structured_blocks):
        raw_block = {
            "page_number": blk["page_number"],
            "type": blk["type"],
            "text": blk.get("text", ""),
        }
        # Pass through definition list metadata
        if blk.get("term") is not None:
            raw_block["term"] = blk["term"]
        if blk.get("definition") is not None:
            raw_block["definition"] = blk["definition"]
        # Pass through heading level if present
        if blk.get("level") is not None:
            raw_block["level"] = blk["level"]
        # Pass through list item metadata
        if blk.get("ordered") is not None:
            raw_block["ordered"] = blk["ordered"]
        raw_blocks.append(raw_block)

    # OCR fallback for scanned/image-based PDFs
    if not pages_data or all(not p["raw_text"].strip() for p in pages_data):
        ocr_result = extract_pdf_ocr(filepath)
        if ocr_result["ocr_info"].get("used"):
            return ocr_result
    return {"blocks": raw_blocks, "metadata": metadata, "pages": len(pages_data)}


def extract_pdf_ocr(filepath: str, dpi: int = 200, max_pages: int = 0) -> dict:
    """OCR fallback for scanned/image-based PDFs using Tesseract.
    
    Args:
        filepath: Path to PDF file
        dpi: Resolution for PDF-to-image conversion (default 200)
        max_pages: Max pages to OCR (0 = all pages). Large PDFs may be slow.
    
    Returns:
        dict with blocks, metadata, pages, and ocr_info
    """
    if not OCR_AVAILABLE:
        return {"blocks": [], "metadata": {}, "pages": 0, "ocr_info": {"used": False, "reason": "OCR libraries not installed"}}
    
    blocks = []
    metadata = {}
    ocr_info = {"used": True, "engine": "tesseract", "version": "", "dpi": dpi, "max_pages": max_pages, "pages_ocred": 0}
    
    try:
        ocr_info["version"] = str(pytesseract.get_tesseract_version())
    except Exception:
        pass
    
    try:
        # Convert PDF pages to images
        kwargs = {"dpi": dpi, "fmt": "png"}
        if max_pages > 0:
            kwargs["first_page"] = 1
            kwargs["last_page"] = max_pages
        
        images = convert_from_path(filepath, **kwargs)
        total_pages = len(images)
        
        for i, image in enumerate(images, start=1):
            # Run OCR on each page
            page_text = pytesseract.image_to_string(image, lang="eng")
            page_text = normalize_whitespace(normalize_unicode(page_text))
            
            if page_text.strip():
                # Try to infer structure from OCR text
                page_blocks = _infer_structure_from_ocr(page_text, page_number=i)
                blocks.extend(page_blocks)
            
            ocr_info["pages_ocred"] = i
        
        metadata["_ocr"] = ocr_info
        return {"blocks": blocks, "metadata": metadata, "pages": total_pages, "ocr_info": ocr_info}
    
    except Exception as e:
        ocr_info["used"] = False
        ocr_info["reason"] = str(e)
        metadata["_ocr_error"] = str(e)
        return {"blocks": [], "metadata": metadata, "pages": 0, "ocr_info": ocr_info}


def _infer_structure_from_ocr(text: str, page_number: int) -> list:
    """Infer document structure from raw OCR text."""
    lines = text.split("\n")
    blocks = []
    current_para = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_para:
                para_text = " ".join(current_para)
                # Check if it looks like a heading
                if _looks_like_heading(para_text):
                    level = _estimate_heading_level(para_text)
                    blocks.append({
                        "type": "heading",
                        "text": para_text,
                        "level": level,
                        "source_location": "ocr inferred",
                        "confidence": 0.75,
                    })
                else:
                    blocks.append({
                        "type": "paragraph",
                        "text": para_text,
                        "source_location": "ocr inferred",
                        "confidence": 0.80,
                    })
                current_para = []
            continue
        
        # Check for list items
        list_match = re.match(r"^(\s*)([-*•]|\d+[.):])\s+(.*)", stripped)
        if list_match and not current_para:
            items = [{"text": list_match.group(3), "level": len(list_match.group(1)) // 2}]
            blocks.append({
                "type": "list",
                "ordered": list_match.group(2)[0].isdigit(),
                "items": items,
                "source_location": "ocr inferred",
                "confidence": 0.75,
            })
            continue
        
        current_para.append(stripped)
    
    # Flush remaining paragraph
    if current_para:
        para_text = " ".join(current_para)
        if _looks_like_heading(para_text):
            level = _estimate_heading_level(para_text)
            blocks.append({
                "type": "heading",
                "text": para_text,
                "level": level,
                "source_location": "ocr inferred",
                "confidence": 0.75,
            })
        else:
            blocks.append({
                "type": "paragraph",
                "text": para_text,
                "source_location": "ocr inferred",
                "confidence": 0.80,
            })
    
    return blocks


def _looks_like_heading(text: str) -> bool:
    """Heuristic: does this text look like a heading?"""
    if not text:
        return False
    # Short text, mostly uppercase, or no sentence-ending punctuation
    words = text.split()
    if len(words) <= 6 and len(text) < 80:
        if text.isupper():
            return True
        if not any(c in text for c in ".!?;"):
            return True
    return False


def _estimate_heading_level(text: str) -> int:
    """Estimate heading level from text characteristics."""
    if text.isupper() and len(text) < 40:
        return 1
    if text.isupper() and len(text) < 60:
        return 2
    if len(text.split()) <= 3:
        return 3
    return 2


# ---------------------------------------------------------------------------
# EPUB extraction
# ---------------------------------------------------------------------------
def extract_epub(filepath: str) -> dict:
    chapters = []
    metadata = {}
    try:
        with zipfile.ZipFile(filepath, "r") as z:
            nl = z.namelist()
            opf_path = None
            if "META-INF/container.xml" in nl:
                container = z.read("META-INF/container.xml").decode("utf-8", errors="replace")
                if etree:
                    root = etree.fromstring(container.encode("utf-8"))
                    ns = {"c": "urn:oasis:names:tc:opendocument:xmlns:container"}
                    rf = root.find(".//c:rootfile", ns)
                    if rf is not None:
                        opf_path = rf.get("full-path")
                else:
                    m = re.search(r'full-path="([^"]+)"', container)
                    if m:
                        opf_path = m.group(1)
            if not opf_path:
                opf_candidates = [n for n in nl if n.endswith(".opf")]
                if opf_candidates:
                    opf_path = opf_candidates[0]
            if opf_path and opf_path in nl:
                opf_xml = z.read(opf_path).decode("utf-8", errors="replace")
                opf_root = etree.fromstring(opf_xml.encode("utf-8")) if etree else None
                if opf_root is not None:
                    ns = {"dc": "http://purl.org/dc/elements/1.1/", "opf": "http://www.idpf.org/2007/opf"}
                    for tag in ("title", "creator", "language", "publisher", "date", "identifier", "description"):
                        el = opf_root.find(f".//dc:{tag}", ns)
                        metadata[tag] = el.text if el is not None else None
                    manifest = {}
                    me = opf_root.find(".//opf:manifest", ns)
                    if me is not None:
                        for item in me.findall("opf:item", ns):
                            manifest[item.get("id")] = {"href": item.get("href"), "media-type": item.get("media-type")}
                    spine = []
                    se = opf_root.find(".//opf:spine", ns)
                    if se is not None:
                        for itemref in se.findall("opf:itemref", ns):
                            spine.append(itemref.get("idref"))
                    opf_dir = str(Path(opf_path).parent) if "/" in opf_path else ""
                    for item_id in spine:
                        if item_id not in manifest:
                            continue
                        href = manifest[item_id]["href"]
                        chapter_path = f"{opf_dir}/{href}".replace("//", "/") if opf_dir else href
                        if chapter_path not in nl:
                            continue
                        html = z.read(chapter_path).decode("utf-8", errors="replace")
                        chapters.append(parse_html_chapter(html, chapter_path, href))
    except Exception as e:
        metadata["_extraction_error"] = str(e)
    return {"chapters": chapters, "metadata": metadata}


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------
def _extract_docx_core_props_from_xml(filepath: str) -> dict:
    """Read DOCX core properties from docProps/core.xml when python-docx is unavailable."""
    props = {}
    try:
        with zipfile.ZipFile(filepath, "r") as z:
            if "docProps/core.xml" not in z.namelist():
                return props
            xml = z.read("docProps/core.xml")
            if etree is None:
                return props
            root = etree.fromstring(xml)
            ns = {
                "dc": "http://purl.org/dc/elements/1.1/",
                "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
                "dcterms": "http://purl.org/dc/terms/",
            }
            # Helper to get text from first matching element
            def _text(tag):
                el = root.find(f".//{tag}", ns)
                return el.text if el is not None else None
            props["title"] = _text("dc:title")
            props["author"] = _text("dc:creator")
            props["subject"] = _text("dc:subject")
            props["language"] = _text("dc:language")
            created_el = root.find(".//dcterms:created", ns)
            props["created"] = created_el.text if created_el is not None else None
            modified_el = root.find(".//dcterms:modified", ns)
            props["modified"] = modified_el.text if modified_el is not None else None
    except Exception:
        pass
    return props


def _extract_docx_xml_fallback(filepath: str) -> tuple[list[dict], list[list[list[str]]]]:
    """Fallback manual XML parsing for DOCX when python-docx fails.
    Returns (paragraphs, tables).
    """
    paragraphs = []
    tables = []
    try:
        with zipfile.ZipFile(filepath, "r") as z:
            if "word/document.xml" not in z.namelist():
                return paragraphs, tables
            xml = z.read("word/document.xml")
            if etree is None:
                return paragraphs, tables
            root = etree.fromstring(xml)
    except Exception:
        return paragraphs, tables

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    body = root.find(".//w:body", ns)
    if body is None:
        return paragraphs, tables

    # Iterate over body children in order to preserve document order
    for child in body:
        tag = etree.QName(child).localname
        if tag == "p":
            # Extract text from all <w:t> elements
            texts = [t.text for t in child.findall(".//w:t", ns) if t.text]
            text = "".join(texts).strip()
            if not text:
                continue
            # Detect heading style
            pPr = child.find("w:pPr", ns)
            heading_level = None
            style_name = ""
            if pPr is not None:
                pStyle = pPr.find("w:pStyle", ns)
                if pStyle is not None:
                    style_val = pStyle.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                    if style_val:
                        style_name = style_val.lower()
                        m = re.search(r"heading\s*(\d+)", style_name, re.IGNORECASE)
                        if m:
                            heading_level = int(m.group(1))
            para_data = {"text": text, "heading_level": heading_level, "style": style_name}
            paragraphs.append(para_data)
        elif tag == "tbl":
            rows = []
            for tr in child.findall("w:tr", ns):
                row = []
                for tc in tr.findall("w:tc", ns):
                    cell_texts = [t.text for t in tc.findall(".//w:t", ns) if t.text]
                    row.append("".join(cell_texts).strip())
                if row:
                    rows.append(row)
            if rows:
                tables.append(rows)
    return paragraphs, tables


def extract_docx(filepath: str) -> dict:
    paragraphs = []
    tables = []
    footnotes = []
    metadata = {}
    core_props = {}
    used_fallback = False
    if Document:
        try:
            doc = Document(filepath)
            core_props = {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
                "language": doc.core_properties.language,
            }
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name or "").lower() if para.style else ""
                heading_level = None
                if "heading" in style_name:
                    try:
                        heading_level = int(re.search(r"\d+", style_name).group())
                    except Exception:
                        heading_level = 1
                para_data = {"text": text, "heading_level": heading_level, "style": style_name}
                paragraphs.append(para_data)
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                tables.append(rows)
        except Exception as e:
            core_props["_error"] = str(e)
    # Fallback: if python-docx failed or is missing, use manual XML parsing
    if not paragraphs and not tables and etree:
        paragraphs, tables = _extract_docx_xml_fallback(filepath)
        used_fallback = True
        # Attempt to read core properties from XML if python-docx didn't set them
        if not core_props or all(v is None for k, v in core_props.items() if not k.startswith("_")):
            core_props.update(_extract_docx_core_props_from_xml(filepath))
    # Footnotes via manual XML
    try:
        with zipfile.ZipFile(filepath, "r") as z:
            if "word/footnotes.xml" in z.namelist():
                xml = z.read("word/footnotes.xml")
                if etree:
                    root = etree.fromstring(xml)
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    for fn in root.findall("w:footnote", ns):
                        fn_id = fn.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
                        if fn_id in ("0", "-1"):
                            continue
                        texts = [t.text for t in fn.findall(".//w:t", ns) if t.text]
                        if texts:
                            footnotes.append({"id": fn_id, "text": " ".join(texts)})
    except Exception:
        pass
    metadata.update(core_props)
    metadata["_extractor"] = "lxml-xml-fallback" if used_fallback else "python-docx"
    return {"paragraphs": paragraphs, "tables": tables, "footnotes": footnotes, "metadata": metadata}


# ---------------------------------------------------------------------------
# Code file handling
# ---------------------------------------------------------------------------
CODE_EXTENSION_MAP = {
    ".py": "python", ".js": "javascript", ".ts": "typescript", ".sh": "bash",
    ".rb": "ruby", ".go": "go", ".rs": "rust", ".java": "java",
    ".cpp": "cpp", ".c": "c", ".h": "c", ".hpp": "cpp", ".cc": "cpp",
    ".php": "php", ".swift": "swift", ".kt": "kotlin", ".scala": "scala",
    ".r": "r", ".m": "objective-c", ".pl": "perl", ".lua": "lua",
    ".sql": "sql", ".json": "json", ".yaml": "yaml", ".yml": "yaml",
    ".xml": "xml", ".css": "css", ".scss": "scss", ".sass": "sass",
    ".html": "html", ".htm": "html", ".bash": "bash", ".zsh": "zsh",
    ".ps1": "powershell", ".bat": "batch", ".cmd": "batch",
    ".dockerfile": "dockerfile", ".makefile": "makefile", ".mk": "makefile",
    ".graphql": "graphql", ".gql": "graphql", ".proto": "protobuf",
    ".toml": "toml", ".ini": "ini", ".cfg": "ini", ".properties": "properties",
    ".gradle": "groovy", ".groovy": "groovy", ".clj": "clojure",
    ".erl": "erlang", ".ex": "elixir", ".exs": "elixir", ".elm": "elm",
    ".hs": "haskell", ".lhs": "haskell", ".nim": "nim", ".dart": "dart",
    ".v": "v", ".fs": "fsharp", ".fsx": "fsharp", ".ml": "ocaml",
    ".mli": "ocaml", ".pas": "pascal", ".dpr": "pascal",
    ".cr": "crystal", ".jl": "julia", ".tcl": "tcl", ".ada": "ada",
    ".adb": "ada", ".ads": "ada", ".coffee": "coffeescript",
    ".litcoffee": "coffeescript", ".jsx": "javascript", ".tsx": "typescript",
    ".vue": "vue", ".svelte": "svelte", ".sol": "solidity",
    ".vy": "vyper", ".cairo": "cairo", ".move": "move",
    ".rsx": "rust", ".rlib": "rust", ".purs": "purescript",
}

SHEBANG_PATTERN = re.compile(r"^#!\s*/usr/bin/env\s+(\S+)|^#!\s*/usr/bin/\S+|^#!\s*/bin/\S+|^#!\s*/sbin/\S+")
FENCED_CODE_RE = re.compile(r"^```(\w+)")


def _detect_language_from_extension(filepath: str) -> str | None:
    ext = Path(filepath).suffix.lower()
    return CODE_EXTENSION_MAP.get(ext)


def _looks_like_code_content(text: str) -> tuple[bool, str | None]:
    """Check if plain text content appears to be code.
    Returns (is_code, detected_language).
    """
    lines = text.split("\n")
    # Check for shebang in first line
    if lines:
        shebang_match = SHEBANG_PATTERN.match(lines[0])
        if shebang_match:
            interpreter = shebang_match.group(1)
            lang_map = {
                "python": "python", "python3": "python", "python2": "python",
                "node": "javascript", "bash": "bash", "sh": "bash", "zsh": "zsh",
                "ruby": "ruby", "perl": "perl", "lua": "lua", "php": "php",
                "r": "r", "julia": "julia", "wish": "tcl",
            }
            return True, lang_map.get(interpreter, interpreter)
    # Check if the very first line is a fenced code block opener.
    # This indicates the entire file is a single code snippet wrapped in a fence
    # (common for .txt files that are actually code exports).
    first_line = lines[0] if lines else ""
    fenced_match = FENCED_CODE_RE.match(first_line)
    if fenced_match:
        return True, fenced_match.group(1) or None
    return False, None


def _extract_code_metadata(text: str, language: str | None) -> dict:
    """Extract function names, class names, and import statements from code text."""
    functions = []
    classes = []
    imports = []
    # Function definitions (Python, JS, TS, Go, Rust, Ruby, PHP, Swift, Kotlin, Java, C-family)
    func_pattern = re.compile(
        r"^[ \t]*(?:async\s+)?(?:def|function|func|fn|void|int|float|double|bool|string|"
        r"var|let|const|public\s+static|private\s+static|protected\s+static|static)?\s*"
        r"([a-zA-Z_][a-zA-Z0-9_]*)\s*\([^)]*\)\s*(?:->\s*\S+\s*)?[{:;=]",
        re.MULTILINE,
    )
    for m in func_pattern.finditer(text):
        name = m.group(1)
        if name not in ("if", "for", "while", "switch", "catch", "else", "elif", "return", "import", "from", "class"):
            functions.append(name)
    # Class definitions
    class_pattern = re.compile(
        r"^[ \t]*(?:class|struct|interface|trait|protocol|enum)\s+([a-zA-Z_][a-zA-Z0-9_]*)",
        re.MULTILINE,
    )
    for m in class_pattern.finditer(text):
        classes.append(m.group(1))
    # Import / include / require / using patterns
    import_patterns = [
        re.compile(r"^[ \t]*(?:import|from)\s+([a-zA-Z_][a-zA-Z0-9_.]*)", re.MULTILINE),  # Python, JS, TS
        re.compile(r"^[ \t]*#include\s+[<\"]([^>\"]+)[>\"]", re.MULTILINE),  # C/C++
        re.compile(r"^[ \t]*(?:require|require_once|include|include_once)\s*\(?['\"]([^'\"]+)['\"]\)?", re.MULTILINE),  # PHP, Ruby
        re.compile(r"^[ \t]*(?:using)\s+([a-zA-Z_][a-zA-Z0-9_.]*);", re.MULTILINE),  # C#, Java, Kotlin, Swift, Go
        re.compile(r"^[ \t]*(?:extern\s+crate|use)\s+([a-zA-Z_][a-zA-Z0-9_:]*);", re.MULTILINE),  # Rust
        re.compile(r"^[ \t]*(?:package)\s+([a-zA-Z_][a-zA-Z0-9_.]*);", re.MULTILINE),  # Java, Kotlin
        re.compile(r"^[ \t]*(?:module)\s+([a-zA-Z_][a-zA-Z0-9_.]*)", re.MULTILINE),  # Ruby, JS
        re.compile(r"^[ \t]*(?:library|require)\s*\(?['\"]([^'\"]+)['\"]\)?", re.MULTILINE),  # R
    ]
    for pat in import_patterns:
        for m in pat.finditer(text):
            imports.append(m.group(1).strip())
    # Deduplicate while preserving order
    seen = set()
    unique_functions = []
    for f in functions:
        if f not in seen:
            seen.add(f)
            unique_functions.append(f)
    seen = set()
    unique_classes = []
    for c in classes:
        if c not in seen:
            seen.add(c)
            unique_classes.append(c)
    seen = set()
    unique_imports = []
    for imp in imports:
        if imp not in seen:
            seen.add(imp)
            unique_imports.append(imp)
    return {
        "functions": unique_functions,
        "classes": unique_classes,
        "imports": unique_imports,
        "line_count": len(text.split("\n")),
    }


# ---------------------------------------------------------------------------
# TXT / Markdown extraction
# ---------------------------------------------------------------------------
def extract_text_file(filepath: str) -> dict:
    with open(filepath, "rb") as f:
        raw = f.read()
    encoding = "utf-8-sig" if raw.startswith(b"\xef\xbb\xbf") else "utf-8"
    try:
        text = raw.decode(encoding)
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="replace")
    # NOTE: We do NOT normalize here; normalization happens in extract_document()
    # so that text_original can be preserved.
    frontmatter = {}
    body = text
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3 and yaml:
            try:
                frontmatter = yaml.safe_load(parts[1])
                body = parts[2].strip()
            except Exception:
                pass

    # --- CODE FILE DETECTION ---
    # 1. Extension-based detection
    detected_lang = _detect_language_from_extension(filepath)
    is_code = detected_lang is not None

    # 2. Content-based detection for .txt and other plain-text files
    if not is_code:
        is_code, detected_lang = _looks_like_code_content(body)

    if is_code:
        meta = _extract_code_metadata(body, detected_lang)
        filename = os.path.basename(filepath)
        block = {
            "type": "code_block",
            "text": body,
            "language": detected_lang,
            "filename": filename,
            "line_count": meta["line_count"],
            "functions": meta["functions"],
            "classes": meta["classes"],
            "imports": meta["imports"],
        }
        return {"blocks": [block], "frontmatter": frontmatter, "metadata": frontmatter if isinstance(frontmatter, dict) else {}}

    is_md = filepath.lower().endswith((".md", ".markdown", ".mdown"))
    if not is_md:
        indicators = ("# ", "## ", "```", "| ", "- [")
        if any(any(line.startswith(i) for i in indicators) for line in body.split("\n")[:30]):
            is_md = True
    if is_md:
        blocks = parse_markdown(body)
    else:
        blocks = parse_plain_text(body)
    return {"blocks": blocks, "frontmatter": frontmatter, "metadata": frontmatter if isinstance(frontmatter, dict) else {}}


def parse_markdown(text: str) -> list[dict]:
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.lstrip()
        # ATX heading
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            blocks.append({"type": "heading", "level": min(level, 6), "text": stripped.lstrip("#").strip()})
            i += 1
            continue
        # Setext h1
        if stripped and i + 1 < len(lines) and lines[i + 1].rstrip() and all(c == "=" for c in lines[i + 1].rstrip()):
            blocks.append({"type": "heading", "level": 1, "text": stripped})
            i += 2
            continue
        # Setext h2
        if stripped and i + 1 < len(lines) and lines[i + 1].rstrip() and all(c == "-" for c in lines[i + 1].rstrip()):
            blocks.append({"type": "heading", "level": 2, "text": stripped})
            i += 2
            continue
        # Fenced code block
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            blocks.append({"type": "code_block", "language": lang or None, "text": "\n".join(code)})
            i += 1
            continue
        # GFM table
        if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
            if re.match(r"^\|?[\s\-:]+(\|[\s\-:]+)*\|?\s*$", lines[i + 1].strip()):
                rows = [[c.strip() for c in line.split("|") if c.strip() or c == ""]]
                rows = [[c for c in row if c] for row in rows]
                i += 2
                while i < len(lines) and "|" in lines[i]:
                    cells = [c.strip() for c in lines[i].split("|") if c.strip() or c == ""]
                    cells = [c for c in cells if c]
                    if cells:
                        rows.append(cells)
                    i += 1
                blocks.append({"type": "table", "rows": rows})
                continue
        # List
        m = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line)
        if m:
            items = [{"text": m.group(3), "level": len(m.group(1)) // 2}]
            i += 1
            while i < len(lines):
                nm = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", lines[i])
                if nm:
                    items.append({"text": nm.group(3), "level": len(nm.group(1)) // 2})
                    i += 1
                elif lines[i].strip() == "":
                    break
                elif lines[i].startswith("  ") or lines[i].startswith("\t"):
                    items[-1]["text"] += " " + lines[i].strip()
                    i += 1
                else:
                    break
            blocks.append({"type": "list", "ordered": m.group(2)[0].isdigit(), "items": items})
            continue
        # Blockquote
        if stripped.startswith(">"):
            quote = []
            while i < len(lines) and (lines[i].lstrip().startswith(">") or lines[i].strip() == ""):
                quote.append(lines[i].lstrip().lstrip(">").strip())
                i += 1
            blocks.append({"type": "quote", "text": " ".join(q for q in quote if q)})
            continue
        # Horizontal rule
        if stripped in ("---", "***", "___", "- - -", "* * *"):
            blocks.append({"type": "divider"})
            i += 1
            continue
        # Paragraph
        if stripped:
            para = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip():
                para.append(lines[i].strip())
                i += 1
            blocks.append({"type": "paragraph", "text": " ".join(para)})
            continue
        i += 1
    return blocks


def parse_plain_text(text: str) -> list[dict]:
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.lstrip()
        if not stripped:
            i += 1
            continue
        if stripped.isupper() and 3 < len(stripped) < 100:
            if i + 1 < len(lines) and not lines[i + 1].strip():
                blocks.append({"type": "heading", "level": 1, "text": stripped.title()})
                i += 1
                continue
        if i + 1 < len(lines):
            nxt = lines[i + 1].rstrip()
            if nxt and all(c == "=" for c in nxt) and len(nxt) >= len(stripped):
                blocks.append({"type": "heading", "level": 1, "text": stripped})
                i += 2
                continue
            if nxt and all(c == "-" for c in nxt) and len(nxt) >= len(stripped):
                blocks.append({"type": "heading", "level": 2, "text": stripped})
                i += 2
                continue
        m = re.match(r"^(\s*)([-*•]|\d+[.):])\s+(.*)", line)
        if m:
            items = [{"text": m.group(3), "level": len(m.group(1)) // 2}]
            i += 1
            while i < len(lines):
                nm = re.match(r"^(\s*)([-*•]|\d+[.):])\s+(.*)", lines[i])
                if nm:
                    items.append({"text": nm.group(3), "level": len(nm.group(1)) // 2})
                    i += 1
                elif lines[i].strip() == "":
                    break
                else:
                    break
            blocks.append({"type": "list", "ordered": m.group(2)[0].isdigit(), "items": items})
            continue
        para = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip():
            # Stop paragraph accumulation if next line looks like a list item or heading
            next_stripped = lines[i].strip()
            if re.match(r"^(\s*)([-*•]|\d+[.):])\s+(.*)", next_stripped):
                break
            if next_stripped.isupper() and 3 < len(next_stripped) < 100:
                break
            if i + 1 < len(lines):
                nxt = lines[i + 1].rstrip()
                if nxt and all(c == "=" for c in nxt) and len(nxt) >= len(next_stripped):
                    break
                if nxt and all(c == "-" for c in nxt) and len(nxt) >= len(next_stripped):
                    break
            para.append(next_stripped)
            i += 1
        blocks.append({"type": "paragraph", "text": "\n".join(para)})
    return blocks


# ---------------------------------------------------------------------------
# HTML extraction
# ---------------------------------------------------------------------------

BLOCK_TAGS = (
    "div", "p", "h1", "h2", "h3", "h4", "h5", "h6",
    "ul", "ol", "li", "table", "tr", "td", "th",
    "pre", "blockquote", "figure", "figcaption",
    "dl", "dt", "dd",
    "section", "article", "aside", "nav", "header", "footer", "main",
    "form", "hr", "br"
)


def _rich_text(elem):
    """Extract text preserving inline semantics."""
    if elem is None:
        return ""
    if isinstance(elem, str):
        return elem

    parts = []
    for child in elem.children:
        if isinstance(child, str):
            # Skip HTML comments (BeautifulSoup Comment is a str subclass)
            if type(child).__name__ == "Comment":
                continue
            parts.append(child)
        elif child.name == "br":
            parts.append("\n")
        elif child.name == "a":
            href = child.get("href", "")
            text = _rich_text(child)
            parts.append(text)
            # Append URL in brackets for ML context
            if href and not href.startswith("#"):
                parts.append(f" [{href}]")
        elif child.name == "code":
            text = _rich_text(child)
            if "\n" not in text:
                parts.append(f"`{text}`")
            else:
                parts.append(text)
        elif child.name == "strong":
            parts.append(f"**{_rich_text(child)}**")
        elif child.name == "em":
            parts.append(f"*{_rich_text(child)}*")
        elif child.name == "mark":
            parts.append(f"=={_rich_text(child)}==")
        elif child.name in ("s", "del", "strike"):
            parts.append(f"~~{_rich_text(child)}~~")
        elif child.name == "abbr":
            title = child.get("title", "")
            text = _rich_text(child)
            parts.append(text)
            if title:
                parts.append(f" ({title})")
        elif child.name == "sub":
            parts.append(_rich_text(child))
        elif child.name == "sup":
            parts.append(_rich_text(child))
        elif child.name == "span":
            parts.append(_rich_text(child))
        elif child.name in ("img",):
            alt = child.get("alt", "")
            if alt:
                parts.append(f"[Image: {alt}]")
        elif child.name in ("input", "button", "select", "textarea"):
            pass  # Skip form elements
        elif child.name in BLOCK_TAGS:
            # Block tags inside inline context - treat as text
            parts.append(_rich_text(child))
        else:
            parts.append(_rich_text(child))

    return "".join(parts)


def extract_html(filepath: str) -> dict:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    return parse_html_chapter(raw, filepath, "")


def parse_html_chapter(html: str, full_path: str = "", href: str = "") -> dict:
    if BeautifulSoup is None:
        return {"blocks": [{"type": "paragraph", "text": normalize_whitespace(normalize_unicode(html))}],
                "metadata": {}}
    soup = BeautifulSoup(html, "html.parser")
    # Strip noise elements
    for tag in soup(["script", "style", "nav", "header", "footer", "aside",
                     "noscript", "iframe", "svg", "canvas", "form", "input",
                     "button", "select", "textarea", "template", "dialog",
                     "address"]):
        tag.decompose()
    metadata = {}
    if soup.title:
        metadata["title"] = soup.title.get_text(strip=True)
    # Enhanced meta tag extraction
    for meta in soup.find_all("meta"):
        prop = meta.get("property", "").lower()
        name = meta.get("name", "").lower()
        content = meta.get("content")
        if content:
            if prop in ("og:title", "og:description", "og:url", "og:type", "og:site_name"):
                metadata[prop.replace(":", "_")] = content
            if name in ("description", "author", "keywords", "language", "twitter:title", "twitter:description"):
                metadata[name.replace(":", "_")] = content
    root = soup.find("main") or soup.find("article") or soup.find("body") or soup
    blocks = []
    for elem in root.find_all(recursive=False):
        block = _convert_html_elem(elem)
        if block:
            if block.get("type") == "container":
                blocks.extend(block.get("blocks", []))
            else:
                blocks.append(block)
    # Language detection enhancement
    all_text = " ".join(b.get("text", "") for b in blocks)
    # Detect Japanese (Hiragana/Katakana/CJK)
    if re.search(r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]', all_text):
        metadata["languages"] = ["en", "ja"]
        metadata["multilingual"] = True
    return {"blocks": blocks, "metadata": metadata, "href": href, "full_path": full_path}


def _convert_html_elem(elem):
    tag = elem.name
    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        return {"type": "heading", "level": int(tag[1]), "text": normalize_whitespace(_rich_text(elem))}
    if tag == "p":
        return {"type": "paragraph", "text": normalize_whitespace(_rich_text(elem))}
    if tag in ("ul", "ol"):
        items = []
        for li in elem.find_all("li", recursive=False):
            # Check for nested lists
            nested_lists = li.find_all(["ul", "ol"], recursive=False)
            text = normalize_whitespace(_rich_text(li))
            # Remove nested list text from the text
            for nl in nested_lists:
                nl_text = normalize_whitespace(_rich_text(nl))
                if nl_text in text:
                    text = text.replace(nl_text, "").strip()
            item = {"text": text}
            if nested_lists:
                item["nested"] = [_convert_html_elem(nl) for nl in nested_lists]
            items.append(item)
        return {"type": "list", "ordered": tag == "ol", "items": items}
    if tag == "table":
        rows = []
        for tr in elem.find_all("tr"):
            row = [normalize_whitespace(_rich_text(td)) for td in tr.find_all(["td", "th"])]
            if row:
                rows.append(row)
        return {"type": "table", "rows": rows}
    if tag in ("pre", "code"):
        code = elem.find("code") if tag == "pre" and elem.find("code") else elem
        classes = code.get("class", []) if hasattr(code, "get") else []
        lang = None
        for c in classes:
            if isinstance(c, str) and (c.startswith("language-") or c.startswith("lang-")):
                lang = c.split("-", 1)[1]
                break
            if isinstance(c, str) and c in ("python", "javascript", "java", "cpp", "bash", "sql"):
                lang = c
        return {"type": "code_block", "language": lang, "text": elem.get_text()}
    if tag == "blockquote":
        return {"type": "quote", "text": normalize_whitespace(_rich_text(elem))}
    if tag in ("div", "section", "article", "main", "header", "footer", "aside"):
        inner = []
        for child in elem.find_all(recursive=False):
            b = _convert_html_elem(child)
            if b:
                if b.get("type") == "container":
                    inner.extend(b.get("blocks", []))
                else:
                    inner.append(b)
        if inner:
            return {"type": "container", "blocks": inner}
        return None
    if tag == "img":
        return {"type": "image_placeholder", "src": elem.get("src", ""), "alt": elem.get("alt", "")}
    if tag == "hr":
        return {"type": "divider"}
    if tag == "figure":
        img = elem.find("img")
        caption = elem.find("figcaption")
        if img and caption:
            return {"type": "figure",
                    "image_src": img.get("src", ""),
                    "alt": img.get("alt", ""),
                    "caption": normalize_whitespace(_rich_text(caption))}
        # Otherwise treat as container
        inner = []
        for child in elem.find_all(recursive=False):
            b = _convert_html_elem(child)
            if b:
                if b.get("type") == "container":
                    inner.extend(b.get("blocks", []))
                else:
                    inner.append(b)
        if inner:
            return {"type": "container", "blocks": inner}
        return None
    if tag == "dl":
        items = []
        current_term = None
        for child in elem.find_all(["dt", "dd"], recursive=False):
            if child.name == "dt":
                current_term = normalize_whitespace(_rich_text(child))
            elif child.name == "dd" and current_term:
                definition = normalize_whitespace(_rich_text(child))
                items.append({"term": current_term, "definition": definition})
        return {"type": "definition_list", "items": items}
    return None


# ---------------------------------------------------------------------------
# Structure building
# ---------------------------------------------------------------------------
def build_structure_and_blocks(raw_blocks: list[dict], source_info: dict) -> tuple[list[dict], dict]:
    """Convert raw extracted blocks into canonical doc2ml blocks and structure tree."""
    blocks = []
    headings = []
    for idx, rb in enumerate(raw_blocks):
        chunk_id = f"blk-{idx:03d}"
        btype = rb.get("type", "paragraph")
        # Convert list_item type to paragraph (structure detected but stored as paragraph)
        if btype in ("list_item",):
            btype = "paragraph"
        # Keep definition_list as-is for special handling below
        if btype == "definition_list":
            pass
        text = rb.get("text_plain", rb.get("text", ""))
        text_original = rb.get("text_original", rb.get("text_plain", rb.get("text", ""))) or text
        if btype == "table":
            rows = rb.get("rows", [])
            rows_original = rb.get("rows_original", rows)
            text = "Table\n" + "\n".join("| " + " | ".join(str(c or "") for c in row) + " |" for row in rows if row)
            text_original = "Table\n" + "\n".join("| " + " | ".join(str(c or "") for c in row) + " |" for row in rows_original if row)
        elif btype == "list":
            items = rb.get("items", [])
            if rb.get("ordered"):
                text = "\n".join(f"{n}. {it.get('text', '')}" for n, it in enumerate(items, start=1))
                text_original = "\n".join(f"{n}. {it.get('text_original', it.get('text', ''))}" for n, it in enumerate(items, start=1))
            else:
                text = "\n".join(f"- {it.get('text', '')}" for it in items)
                text_original = "\n".join(f"- {it.get('text_original', it.get('text', ''))}" for it in items)
        elif btype == "divider":
            text = "---"
        elif btype == "image_placeholder":
            text = rb.get("alt", "") or rb.get("src", "")
        page_num = rb.get("page_number") or rb.get("page")
        b = make_block_base(chunk_id, btype, text, page_num, rb.get("source_location", ""),
                            source_info.get("extractor", ""), rb.get("confidence", 0.95))
        b["text_original"] = text_original if text_original else text
        if btype == "heading":
            level = rb.get("level") or _estimate_heading_level_from_text(text)
            b["content"] = {"text": text, "level": level, "numbered": False, "label": None}
            b["semantics"]["heading_level"] = level
            headings.append({"chunk_id": chunk_id, "level": level, "text": text})
        elif btype == "paragraph":
            b["content"] = {"text": text, "inline_elements": [], "sentences": []}
            kvs = extract_key_values(text)
            if kvs:
                b["custom"]["derived"] = {"key_values": kvs}
        elif btype == "code_block":
            # Use text_original to preserve indentation for code
            code_text = text_original if text_original else text
            b["content"] = {
                "code": code_text,
                "language": rb.get("language"),
                "filename": rb.get("filename"),
                "line_numbers": False,
                "lines": [{"line_number": li + 1, "text": ln} for li, ln in enumerate(code_text.split("\n"))],
                "line_count": rb.get("line_count"),
                "functions": rb.get("functions", []),
                "classes": rb.get("classes", []),
                "imports": rb.get("imports", []),
            }
        elif btype == "list":
            b["content"] = {"list_type": "ordered" if rb.get("ordered") else "unordered", "items": rb.get("items", [])}
        elif btype == "table":
            rows = rb.get("rows", [])
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []
            b["content"] = {
                "headers": headers,
                "rows": data_rows,
                "cells": rows,
                "row_count": len(data_rows),
                "header_row_count": 1 if headers else 0,
                "column_count": max(len(r) for r in rows) if rows else 0,
            }
        elif btype == "definition_list":
            b["content"] = {
                "term": rb.get("term", ""),
                "definition": rb.get("definition", ""),
                "items": [{"term": rb.get("term", ""), "definition": rb.get("definition", "")}],
            }
        elif btype == "quote":
            b["content"] = {"text": text, "attribution": None, "source": None, "cite_chunk_id": None}
        elif btype == "image_placeholder":
            b["content"] = {"image_id": f"img-{idx:03d}", "image_uri": rb.get("src", ""), "alt_text": rb.get("alt", ""),
                            "width_px": None, "height_px": None, "format": None}
        elif btype == "divider":
            b["content"] = {}
        else:
            b["content"] = {"raw": rb}
        blocks.append(b)

    # Post-process: merge short label paragraphs into following list blocks
    merged_blocks = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if (b["type"] == "paragraph" and i + 1 < len(blocks) and blocks[i + 1]["type"] == "list"):
            words = b["text_plain"].split()
            stripped = b["text_plain"].rstrip()
            last_char = stripped[-1:] if stripped else ""
            if len(words) <= 5 and last_char not in ".!?;" and last_char == ":":
                next_block = blocks[i + 1]
                next_block["content"]["intro_text"] = b["text_plain"]
                next_block["text_plain"] = b["text_plain"] + "\n" + next_block["text_plain"]
                next_block["text_original"] = b["text_original"] + "\n" + next_block["text_original"]
                next_block["char_count"] = len(next_block["text_plain"])
                next_block["token_count_est"] = estimate_tokens(next_block["text_plain"])
                merged_blocks.append(next_block)
                i += 2
                continue
        merged_blocks.append(b)
        i += 1
    blocks = merged_blocks

    # Link context windows
    for i, b in enumerate(blocks):
        b["context_window"]["prev_chunk_id"] = blocks[i - 1]["chunk_id"] if i > 0 else None
        b["context_window"]["next_chunk_id"] = blocks[i + 1]["chunk_id"] if i + 1 < len(blocks) else None
        # Find nearest preceding heading with proper hierarchy
        parent_heading = None
        if b["type"] == "heading":
            current_level = b["semantics"]["heading_level"]
            for j in range(i - 1, -1, -1):
                if blocks[j]["type"] == "heading" and blocks[j]["semantics"]["heading_level"] < current_level:
                    parent_heading = blocks[j]["chunk_id"]
                    break
        else:
            for j in range(i - 1, -1, -1):
                if blocks[j]["type"] == "heading":
                    parent_heading = blocks[j]["chunk_id"]
                    break
        b["context_window"]["parent_heading_chunk_id"] = parent_heading
        # Preview
        prev_text = blocks[i - 1]["text_plain"][:50] if i > 0 else ""
        next_text = blocks[i + 1]["text_plain"][:50] if i + 1 < len(blocks) else ""
        b["context_window"]["surrounding_text_preview"] = f"{prev_text} | {next_text}".strip(" |")
    # Build structure tree
    structure = build_structure_tree(blocks, headings)
    return blocks, structure


def build_structure_tree(blocks: list[dict], headings: list[dict]) -> dict:
    """Build hierarchical structure from heading list and blocks."""
    root = {"node_id": "root", "node_type": "document", "title": "", "level": 0, "chunk_ids": [], "children": []}
    if not headings:
        # Flat: all blocks under root
        root["chunk_ids"] = [b["chunk_id"] for b in blocks]
        return root
    # Normalize heading levels
    levels = sorted(set(h["level"] for h in headings))
    level_map = {old: new for new, old in enumerate(levels, start=1)}
    for h in headings:
        h["level"] = level_map[h["level"]]
    # Ensure h1 exists
    if headings[0]["level"] > 1:
        title_guess = blocks[0]["text_plain"][:100] if blocks else "Untitled Document"
        root["title"] = title_guess
    # Build nodes
    stack = [root]
    heading_idx = 0
    for i, block in enumerate(blocks):
        if block["type"] == "heading":
            level = headings[heading_idx]["level"]
            node = {"node_id": f"sec-{heading_idx:03d}", "node_type": _level_to_node_type(level),
                    "title": block["text_plain"], "level": level, "chunk_ids": [block["chunk_id"]], "children": []}
            while stack and stack[-1]["level"] >= level:
                stack.pop()
            if stack:
                stack[-1]["children"].append(node)
            else:
                root["children"].append(node)
            stack.append(node)
            block["context_window"]["parent_structure_node_id"] = node["node_id"]
            block["semantics"]["section_role"] = _guess_section_role(block["text_plain"])
            heading_idx += 1
        else:
            if stack:
                stack[-1]["chunk_ids"].append(block["chunk_id"])
                block["context_window"]["parent_structure_node_id"] = stack[-1]["node_id"]
            else:
                root["chunk_ids"].append(block["chunk_id"])
                block["context_window"]["parent_structure_node_id"] = "root"
    return root


def _level_to_node_type(level: int) -> str:
    mapping = {1: "section", 2: "subsection", 3: "subsubsection", 4: "subsubsection", 5: "subsubsection", 6: "subsubsection"}
    return mapping.get(level, "section")


def _guess_section_role(text: str) -> str | None:
    lower = text.lower()
    keywords = {
        "introduction": "introduction", "background": "background", "related work": "background",
        "method": "method", "methods": "method", "methodology": "method", "approach": "method",
        "experiment": "results", "experiments": "results", "results": "results", "evaluation": "results",
        "discussion": "discussion", "conclusion": "conclusion", "conclusions": "conclusion",
        "abstract": "abstract", "references": "reference", "acknowledgment": "metadata",
    }
    for kw, role in keywords.items():
        if kw in lower:
            return role
    return None


def build_ml_index(blocks: list[dict], structure: dict) -> dict:
    """Build ML acceleration index."""
    chunk_id_map = {}
    heading_map = []
    embedding_candidates = []
    chunk_boundaries = []
    _walk_structure(structure, [], chunk_id_map, heading_map)
    for i, b in enumerate(blocks):
        chunk_id_map.setdefault(b["chunk_id"], {"index": i, "structure_path": ["root"]})
        if b["embedding_ready"]:
            embedding_candidates.append(b["chunk_id"])
    # Build chunk boundaries from structure sections
    _walk_boundaries(structure, blocks, chunk_boundaries)
    return {
        "chunk_id_map": chunk_id_map,
        "heading_map": heading_map,
        "embedding_candidates": embedding_candidates,
        "chunk_boundaries": chunk_boundaries,
    }


def _walk_structure(node: dict, path: list[str], chunk_id_map: dict, heading_map: list[dict]):
    current_path = path + [node["node_id"]]
    for cid in node.get("chunk_ids", []):
        chunk_id_map[cid] = {"index": 0, "structure_path": current_path.copy()}
    if node.get("title") and node.get("chunk_ids"):
        heading_map.append({"chunk_id": node["chunk_ids"][0], "heading_text": node["title"], "level": node["level"], "node_id": node["node_id"]})
    for child in node.get("children", []):
        _walk_structure(child, current_path, chunk_id_map, heading_map)


def _walk_boundaries(node: dict, blocks: list[dict], boundaries: list[dict]):
    cids = node.get("chunk_ids", [])
    for child in node.get("children", []):
        child_cids = child.get("chunk_ids", [])
        if child_cids:
            tok_est = sum(b["token_count_est"] for b in blocks if b["chunk_id"] in child_cids)
            boundaries.append({"start_chunk_id": child_cids[0], "end_chunk_id": child_cids[-1],
                                 "boundary_type": "section", "token_count_est": tok_est})
        _walk_boundaries(child, blocks, boundaries)


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------
def extract_document(filepath: str) -> dict:
    start = time.time()
    mime = detect_format(filepath)
    filename = os.path.basename(filepath)
    size = os.path.getsize(filepath)
    doc_id = str(uuid.uuid4())
    extractor = "doc2ml-json"
    extractor_version = "0.6.2"
    pipeline_steps = ["detect", "extract", "normalize", "structure", "index"]
    raw_blocks = []
    meta_extra = {}
    page_count = 0
    if mime == "application/pdf":
        result = extract_pdf(filepath)
        raw_blocks = result["blocks"]
        meta_extra = result.get("metadata", {})
        page_count = result.get("pages", 0)
        extractor = "pdfplumber" if pdfplumber else "pymupdf"
        if "ocr_info" in result:
            meta_extra["_ocr_info"] = result["ocr_info"]
            if result["ocr_info"].get("used"):
                extractor = f"{extractor}+tesseract_ocr"
    elif mime == "application/epub+zip":
        result = extract_epub(filepath)
        chapters = result.get("chapters", [])
        meta_extra = result.get("metadata", {})
        page_count = len(chapters)
        extractor = "ebooklib"
        for ch in chapters:
            raw_blocks.extend(ch.get("blocks", []))
    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        result = extract_docx(filepath)
        meta_extra = result.get("metadata", {})
        paragraphs = result.get("paragraphs", [])
        for p in paragraphs:
            raw_blocks.append({"type": "heading" if p.get("heading_level") else "paragraph",
                               "text": p["text"], "level": p.get("heading_level"),
                               "source_location": "paragraph"})
        for t in result.get("tables", []):
            raw_blocks.append({"type": "table", "rows": t, "source_location": "table"})
        for fn in result.get("footnotes", []):
            raw_blocks.append({"type": "footnote", "text": fn.get("text", ""), "source_location": "footnote"})
        extractor = meta_extra.get("_extractor", "python-docx")
    elif mime in ("text/plain", "text/markdown"):
        result = extract_text_file(filepath)
        raw_blocks = result["blocks"]
        meta_extra = result.get("metadata", {})
        extractor = "native"
    elif mime == "text/html":
        result = extract_html(filepath)
        raw_blocks = result["blocks"]
        meta_extra = result.get("metadata", {})
        extractor = "beautifulsoup"
    else:
        result = extract_text_file(filepath)
        raw_blocks = result["blocks"]
        meta_extra = result.get("metadata", {})
        extractor = "native"
    # Normalize all text, preserving originals for text_original field
    def _normalize_in_place(obj, key="text"):
        if isinstance(obj, dict):
            if key in obj and isinstance(obj[key], str):
                obj[f"{key}_original"] = obj[key]
                obj[key] = normalize_whitespace(normalize_unicode(obj[key]))
            for v in obj.values():
                _normalize_in_place(v, key)
        elif isinstance(obj, list):
            for item in obj:
                _normalize_in_place(item, key)
    for rb in raw_blocks:
        # Preserve table rows before normalizing cells
        if rb.get("type") == "table" and "rows" in rb:
            rb["rows_original"] = [[cell for cell in row] for row in rb["rows"]]
        _normalize_in_place(rb, "text")
        # Normalize table cell strings in rows
        if rb.get("type") == "table" and "rows" in rb:
            for row in rb["rows"]:
                for ci, cell in enumerate(row):
                    if isinstance(cell, str):
                        row[ci] = normalize_whitespace(normalize_unicode(cell))
    source_info = {"page": None, "extractor": extractor}
    blocks, structure = build_structure_and_blocks(raw_blocks, source_info)
    # Update parent heading and structure refs after tree built
    ml_index = build_ml_index(blocks, structure)
    # Extract cross-references from all block text
    cross_references = _extract_cross_references_from_blocks(blocks)
    # Extract dates from first 5000 chars of block text for metadata
    combined_text = " ".join(b.get("text_plain", "") for b in blocks)
    extracted_dates = _extract_dates_from_text(combined_text[:5000])
    # Compute stats
    char_count = sum(b["char_count"] for b in blocks)
    tok_count = sum(b["token_count_est"] for b in blocks)
    word_count = sum(len(b["text_plain"].split()) for b in blocks)
    heading_count = sum(1 for b in blocks if b["type"] == "heading")
    table_count = sum(1 for b in blocks if b["type"] == "table")
    figure_count = sum(1 for b in blocks if b["type"] == "figure_caption")
    fn_count = sum(1 for b in blocks if b["type"] == "footnote")
    duration = int((time.time() - start) * 1000)
    doc = {
        "doc2ml_version": "0.6.2",
        "document_id": doc_id,
        "metadata": {
            "title": meta_extra.get("title") or meta_extra.get("Title") or filename,
            "subtitle": meta_extra.get("subtitle") or meta_extra.get("Subject"),
            "authors": [{"name": a}] if isinstance(a := meta_extra.get("author") or meta_extra.get("creator"), str) else [],
            "source": {
                "uri": filepath,
                "mime_type": mime,
                "filename": filename,
                "checksum_sha256": sha256_file(filepath),
                "file_size_bytes": size,
                "declared_source": meta_extra.get("source") or filepath,
                "source_type": "frontmatter" if meta_extra.get("source") else "filesystem",
            },
            "ingestion": {
                "ingestion_date": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
                "processing_version": "doc2ml-json v0.6.2",
                "extractor": extractor,
                "extractor_version": extractor_version,
                "ingestion_pipeline": pipeline_steps,
                "processing_duration_ms": duration,
            },
            "language": (
                {"detected": "multilingual", "confidence": 0.95, "languages": meta_extra.get("languages", [])}
                if meta_extra.get("multilingual")
                else {"detected": "en", "confidence": 0.99}
            ),
            "statistics": {
                "page_count": page_count,
                "chapter_count": 0,
                "section_count": heading_count,
                "block_count": len(blocks),
                "table_count": table_count,
                "figure_count": figure_count,
                "footnote_count": fn_count,
                "total_char_count": char_count,
                "total_token_count_est": tok_count,
                "total_word_count": word_count,
            },
            "classification": {"doc_type": "unknown", "genre": None, "keywords": [], "topics_ml": []},
            "dates": {"created": meta_extra.get("created") or meta_extra.get("date"),
                      "modified": meta_extra.get("modified"),
                      "published": meta_extra.get("published") or (extracted_dates[0] if extracted_dates else None)},
            "rights": {"license": None, "copyright": None, "open_access": True},
            "_ocr_info": meta_extra.get("_ocr_info"),
        },
        "structure": structure,
        "blocks": blocks,
        "cross_references": cross_references,
        "ml_index": ml_index,
        "custom": {},
    }
    return doc


def main():
    parser = argparse.ArgumentParser(description="Extract structured JSON from any document")
    parser.add_argument("filepath", help="Path to input document")
    parser.add_argument("-o", "--output", help="Output JSON file path (default: {document_id}.doc2ml.json)")
    parser.add_argument("--output-dir", help="Directory to write output JSON")
    args = parser.parse_args()
    try:
        doc = extract_document(args.filepath)
        out_path = args.output
        if not out_path and args.output_dir:
            out_path = os.path.join(args.output_dir, f"{doc['document_id']}.doc2ml.json")
        if not out_path:
            out_path = f"{doc['document_id']}.doc2ml.json"
        os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
        doc = extract_document(args.filepath)
        doc = _serialize_dates(doc)
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(doc, f, indent=2, ensure_ascii=False)
        print(f"Extracted to {out_path} ({doc['metadata']['statistics']['block_count']} blocks, "
              f"{doc['metadata']['statistics']['total_token_count_est']} tokens)")
    except Exception as e:
        print(json.dumps({"error": str(e)}), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
